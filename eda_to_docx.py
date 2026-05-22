# -*- coding: utf-8 -*-
"""המרת EDA_REPORT.md לקובץ Word (.docx) עם עיצוב עברית RTL ומיתוג i24.

קורא את EDA_REPORT.md, מנתח Markdown (כותרות, טבלאות, רשימות, ציטוטים,
טקסט מודגש/קוד) ומפיק EDA_REPORT.docx.

הרצה:  py -3 eda_to_docx.py
פלט:   EDA_REPORT.docx
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# שימוש:  py -3 eda_to_docx.py [SRC.md] [OUT.docx]
# ברירת מחדל (תאימות לאחור): EDA_REPORT.md -> EDA_REPORT.docx
SRC = sys.argv[1] if len(sys.argv) > 1 else "EDA_REPORT.md"
OUT = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(SRC)[0] + ".docx"
FONT = "David"

# ---- צבעי מותג i24 ----
PRIMARY = RGBColor(0x1E, 0x40, 0xAF)    # כחול עמוק
SECONDARY = RGBColor(0x2D, 0x37, 0x48)  # אפור כהה
ACCENT = RGBColor(0xF5, 0x9E, 0x0B)     # ענבר
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)

HEADING_SIZE = {1: 20, 2: 16, 3: 13, 4: 12}


def set_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def style_run(run, size=11, bold=False, italic=False, color=None, mono=False):
    run.font.name = "Consolas" if mono else FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


# ---- ניתוח עיצוב inline: **מודגש** ו-`קוד` ----
def add_inline(paragraph, text, size=11, base_bold=False, base_color=None):
    """מוסיף runs לפסקה תוך פירוק **מודגש** ו-`קוד`."""
    # לכווץ רצף של 3+ כוכביות ל-2 (יש בקובץ ****טקסט**)
    text = re.sub(r'\*{3,}', '**', text)
    # פיצול לפי **bold** — האינדקסים האי-זוגיים מודגשים
    bold_parts = text.split('**')
    for bi, bpart in enumerate(bold_parts):
        if bpart == '':
            continue
        is_bold = base_bold or (bi % 2 == 1)
        # פיצול לפי `code`
        code_parts = bpart.split('`')
        for ci, cpart in enumerate(code_parts):
            if cpart == '':
                continue
            is_code = (ci % 2 == 1)
            run = paragraph.add_run(cpart)
            style_run(run, size=size, bold=is_bold,
                      color=PRIMARY if is_code else base_color,
                      mono=is_code)


def new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    color = PRIMARY if level <= 2 else SECONDARY
    add_inline(p, text, size=HEADING_SIZE.get(level, 12),
               base_bold=True, base_color=color)
    if level == 1:
        # קו תחתון דק לכותרת ראשית
        pPr = p._element.get_or_add_pPr()
        bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1E40AF')
        bdr.append(bottom)
        pPr.append(bdr)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text, size=11)


def add_quote(doc, text):
    p = new_para(doc, space_after=8)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(6)
    # רקע בהיר + פס צבעוני
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFF7E6')
    pPr.append(shd)
    bdr = OxmlElement('w:pBdr')
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '18')
    right.set(qn('w:space'), '8')
    right.set(qn('w:color'), 'F59E0B')
    bdr.append(right)
    pPr.append(bdr)
    add_inline(p, text, size=11, base_color=SECONDARY)
    for r in p.runs:
        r.font.italic = True


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def add_table(doc, rows):
    """rows: רשימת רשימות-מחרוזות. השורה הראשונה היא הכותרת."""
    headers = rows[0]
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.autofit = True

    # גודל גופן יורד ככל שיש יותר עמודות
    fsize = 9 if ncols <= 6 else (8 if ncols <= 10 else 7)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.text = ""
        para = cell.paragraphs[0]
        set_rtl(para)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        add_inline(para, h or " ", size=fsize, base_bold=True,
                   base_color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, '1E40AF')

    for r, row in enumerate(rows[1:]):
        cells = table.add_row().cells
        for i in range(ncols):
            val = row[i] if i < len(row) else ""
            cell = cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            set_rtl(para)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.space_after = Pt(0)
            add_inline(para, val if val else " ", size=fsize)
        if r % 2 == 1:
            for c in cells:
                shade_cell(c, 'EEF2FB')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def split_table_row(line, ncols):
    """מפצל שורת טבלה ל-cells. ממזג עודף תאים (פייפים לא-בורחים) לתא האחרון."""
    raw = line.strip()
    if raw.startswith('|'):
        raw = raw[1:]
    if raw.endswith('|'):
        raw = raw[:-1]
    parts = [c.strip() for c in raw.split('|')]
    if len(parts) > ncols:
        parts = parts[:ncols - 1] + [' | '.join(parts[ncols - 1:])]
    return parts


def is_table_sep(line):
    return bool(re.match(r'^\s*\|?[\s:|-]+\|?\s*$', line)) and '-' in line


# ========================================================================
with open(SRC, encoding='utf-8') as f:
    lines = f.read().split('\n')

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

# גופן ברירת מחדל
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(11)

i = 0
para_buf = []


def flush_para():
    global para_buf
    if not para_buf:
        return
    text = ' '.join(s.strip() for s in para_buf).strip()
    para_buf = []
    if not text:
        return
    # פסקה שכולה נטויה: *טקסט*  →  כיתוב מעומעם
    m = re.fullmatch(r'\*([^*].*?)\*', text)
    if m:
        p = new_para(doc)
        add_inline(p, m.group(1), size=10, base_color=MUTED)
        for r in p.runs:
            r.font.italic = True
        return
    p = new_para(doc)
    add_inline(p, text, size=11, base_color=TEXT)


cover_done = False
quote_buf = []


def flush_quote():
    global quote_buf
    if quote_buf:
        add_quote(doc, ' '.join(quote_buf).strip())
        quote_buf = []


while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # שורה ריקה
    if stripped == '':
        flush_para()
        flush_quote()
        i += 1
        continue

    # קו מפריד
    if re.fullmatch(r'-{3,}', stripped) or re.fullmatch(r'\*{3,}', stripped):
        flush_para()
        flush_quote()
        i += 1
        continue

    # כותרת
    hm = re.match(r'^(#{1,6})\s+(.*)$', stripped)
    if hm:
        flush_para()
        flush_quote()
        level = len(hm.group(1))
        htext = hm.group(2).strip()
        if not cover_done and level == 1:
            # שער
            title = doc.add_paragraph()
            set_rtl(title)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_after = Pt(4)
            add_inline(title, htext, size=30, base_bold=True,
                       base_color=PRIMARY)
            cover_done = True
        else:
            add_heading(doc, htext, min(level, 4))
        i += 1
        continue

    # טבלה: שורה נוכחית מתחילה ב-| והשורה הבאה היא מפריד
    if stripped.startswith('|') and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
        flush_para()
        flush_quote()
        header = split_table_row(lines[i], 999)
        ncols = len(header)
        tbl_rows = [header]
        j = i + 2
        while j < len(lines) and lines[j].strip().startswith('|'):
            if is_table_sep(lines[j]):
                j += 1
                continue
            tbl_rows.append(split_table_row(lines[j], ncols))
            j += 1
        add_table(doc, tbl_rows)
        i = j
        continue

    # ציטוט / callout
    if stripped.startswith('>'):
        flush_para()
        qtext = re.sub(r'^>\s?', '', stripped)
        if qtext.strip():
            quote_buf.append(qtext)
        i += 1
        continue
    else:
        flush_quote()

    # רשימה
    bm = re.match(r'^[-*]\s+(.*)$', stripped)
    if bm:
        flush_para()
        add_bullet(doc, bm.group(1).strip())
        i += 1
        continue

    # פסקה רגילה — לצבור
    para_buf.append(stripped)
    i += 1

flush_para()
flush_quote()

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT) / 1024:.0f} KB")
