# -*- coding: utf-8 -*-
"""Generate a Word (.docx) version of PROJECT_SUMMARY.md.

Uses python-docx with Hebrew RTL formatting + i24 brand styling.

Run: py -3 generate_docx.py
Output: פרוייקט_i24_סיכום.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ---- i24 brand colors ----
PRIMARY = RGBColor(0x1E, 0x40, 0xAF)   # deep blue
SECONDARY = RGBColor(0x2D, 0x37, 0x48)  # dark gray
ACCENT = RGBColor(0xF5, 0x9E, 0x0B)     # amber
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def set_rtl(paragraph):
    """Set RTL direction on paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def set_cell_rtl(cell):
    """Set RTL on table cell."""
    for p in cell.paragraphs:
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(h)
    for run in h.runs:
        run.font.name = "David"
        run.font.size = Pt({1: 22, 2: 16, 3: 14, 4: 12}.get(level, 12))
        run.font.color.rgb = color or PRIMARY
        run.font.bold = True
    return h


def add_para(doc, text, bold=False, color=None, size=11, align="right"):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = {"right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    run = p.add_run(text)
    run.font.name = "David"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.name = "David"
            run.font.size = Pt(11)


def add_table(doc, headers, rows, header_color=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = h
        set_cell_rtl(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "David"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Cell shading
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1E40AF')
        tc_pr.append(shd)

    # Data rows
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = str(val)
            set_cell_rtl(cells[i])
            for run in cells[i].paragraphs[0].runs:
                run.font.name = "David"
                run.font.size = Pt(10)
    return table


def add_callout(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"💡 {text}")
    run.font.name = "David"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = ACCENT
    # Indent
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)


def add_quote(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'"{text}"')
    run.font.name = "David"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MUTED


# ========================================================================
# Build the document
# ========================================================================
doc = Document()

# Set page margins
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ----- Cover -----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(title)
run = title.add_run("📺 פרוייקט i24")
run.font.name = "David"
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = PRIMARY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(subtitle)
run = subtitle.add_run("חיזוי רייטינג טלוויזיה — מ-EDA למוצר חי באוויר")
run.font.name = "David"
run.font.size = Pt(18)
run.font.color.rgb = SECONDARY

byline = doc.add_paragraph()
byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(byline)
run = byline.add_run("אורית עלמה זיו-נר | מאי 2025 – מאי 2026")
run.font.name = "David"
run.font.size = Pt(12)
run.font.color.rgb = MUTED

doc.add_paragraph()
doc.add_paragraph()

# ----- 1. Executive Summary -----
add_heading(doc, "1. תקציר מנהלים", level=1)
add_para(doc, "מערכת חיזוי רייטינג טלוויזיה לערוץ i24 — מ-EDA עד מוצר חי באוויר.")
add_para(doc, "הישגים:", bold=True)
add_bullets(doc, [
    "19 מודלים נבדקו, HistGradientBoosting ניצח (MAE=0.263, R²=0.603)",
    "שיפור 37.8% מעל המודל הנאיבי",
    "אפליקציית Streamlit חיה: https://i24-ratings-orit.streamlit.app",
    "קוד פתוח ב-GitHub: https://github.com/DataScienceOritAlma/i24-ratings-forecast",
    "5 מסכים פעילים: leaderboard, חיזויים, כרטיס תוכנית, השוואה, חיזוי עתידי",
])
add_callout(doc, "תובנה מרכזית: תקרת הביצועים היא drift של אירועים בלתי-צפויים, "
                  "לא בחירת מודל. 19 מודלים שונים לא משפרים יותר מ-2.6% מעל RandomForest.")

doc.add_page_break()

# ----- 2. The Problem -----
add_heading(doc, "2. הבעיה העסקית — לפי מנהל המחקר של i24", level=1)

add_heading(doc, "המטרה המרכזית", level=2)
add_quote(doc, "ניבוי רייטינג של תוכניות. כל שאר הפרמטרים נועדו לסייע. "
                "המטרה המרכזית היא ניבוי רייטינג.")

add_heading(doc, "אופק הזמן", level=2)
add_para(doc, "סוף שנה / כמה חודשים קדימה — לא חיזוי לעוד שבוע, אלא תכנון אסטרטגי "
              "לטווח בינוני-ארוך.")

add_heading(doc, "הערך העסקי — למה כל הארגון תלוי בזה", level=2)
add_table(doc,
    headers=["תחום", "מה מתאפשר"],
    rows=[
        ("תכנון", "בניית תוכניות עבודה מבוססות נתונים"),
        ("כספים", "תחזית הכנסות (מחיר פרסום נגזר מרייטינג)"),
        ("תקציב", "תחזית הוצאות (כמה משאבים להשקיע בהפקה)"),
        ("לוז", "שינויי לוח שידורים, הוספה/הסרת תוכניות, סדר השידור"),
    ],
)
add_quote(doc, "כל הארגון תלוי בנתונים האלה.")

add_heading(doc, "אבולוציה עתידית — מתוכנית-בודדת לרצועות", level=2)
add_para(doc, "ברגע שכלי הניבוי לתוכניות עובד טוב, ניתן להתפתח ל:")
add_bullets(doc, [
    "ניבוי רצועות כלליות (לדוגמה: כל רצועת הפריים-טיים)",
    "אופטימיזציה של פרסומות — כמה דקות פרסום אפשר להכניס לרצועה מבלי לפגוע ברייטינג מעבר לרף",
    "סימולציות תכנון — מה יקרה אם נעביר תוכנית X מ-19:00 ל-21:00",
    "חישובי trade-off — איזון בין הכנסות פרסום ל-engagement",
])

add_heading(doc, "איך זה מתורגם ל-MVP", level=2)
add_bullets(doc, [
    "MVP1 (קיים): ניבוי רייטינג של תוכנית בודדת, אופק עד שנה קדימה",
    "MVP2 (next): ניבוי רצועות (פריים, אחה\"צ) — אגרגציה של חיזויי תוכניות",
    "MVP3 (vision): מחשבון פרסומות — דקות פרסום → צפי ירידה ברייטינג",
])

# ----- 3. The Data -----
add_heading(doc, "3. הדאטה", level=1)
add_table(doc,
    headers=["נתון", "ערך"],
    rows=[
        ("מקור", "לוג שידורים פנימי של i24"),
        ("טווח זמנים", "25/05/2025 → 18/04/2026"),
        ("ימים מכוסים", "329"),
        ("שורות", "10,039"),
        ("תוכניות ייחודיות", "179"),
        ("עמודות גולמיות", "15"),
        ("ערכים חסרים", "0"),
        ("יחס פיצול", "80/20 כרונולוגי (חיתוך 2026-02-08)"),
        ("עמודת היעד", "רייטינג (0.0-5.0+)"),
    ],
)

# ----- 4. EDA Findings -----
doc.add_page_break()
add_heading(doc, "4. ממצאי EDA — 5 תובנות מרכזיות", level=1)

add_heading(doc, "4.1 פריים-טיים אצל i24 = 17:00-19:00", level=2)
add_para(doc, "ערוץ חדשות — רוב הצפייה לפני מהדורת הערב של ערוצי הביידור (20:00).")

add_heading(doc, "4.2 47.6% מהשידורים חזרות", level=2)
add_para(doc, "לקט נחשב גם כחזרה. צריך להבדיל לגמרי בין שידור חי לחזרה.")

add_heading(doc, "4.3 אירועים ביטחוניים מקפיצים רייטינג פי 4-10", level=2)
add_para(doc, "מבצע שאגת הארי (28/02/2026): תוכניות 'מיוחד-מתקפה באיראן' "
              "קיבלו רייטינג 4.7-5.0 כשהממוצע הוא 0.4.")

add_heading(doc, "4.4 Drift אמיתי בין Train ל-Test", level=2)
add_para(doc, "Train mean: 0.412 | Test mean: 0.558 (פער של 35%). "
              "אירועי 2026-Q1 הזיזו את ההתפלגות.")

add_heading(doc, "4.5 הפאנל הנדגם גדל במהלך השנה", level=2)
add_para(doc, "מ-65% ל-90% בעלי-מקלטים. הוספתי reception_pct כ-feature.")

# ----- 5. Feature Engineering -----
doc.add_page_break()
add_heading(doc, "5. הנדסת פיצ'רים — 19 פיצ'רים מהונדסים", level=1)
add_table(doc,
    headers=["משפחה", "פיצ'רים"],
    rows=[
        ("זמן (4)", "שעת התחלה, חלקי-יום, חודש, שבוע-בשנה"),
        ("חזרות (3)", "is_rerun, שם תוכנית_מקור, סטטוס תוכנית"),
        ("ביצועים (3)", "ממוצע מתחרים, יתרון מול מתחרים, HUT proxy"),
        ("נירמול (2)", "reception_pct, רייטינג מותאם"),
        ("אירועים (7)", "תג_עונה, תג_חג, תג_ביטחוני, אירוע_מיוחד, יום_חג, יום_ביטחוני, שבת"),
    ],
)

add_heading(doc, "Lag Features (במידול)", level=2)
add_bullets(doc, [
    "lag_program_mean — ממוצע היסטורי של אותה תוכנית",
    "lag_slot_mean — ממוצע היסטורי באותה רצועה (יום × שעה)",
    "lag_competitors_avg_slot — ממוצע מתחרים באותה רצועה",
])
add_callout(doc, "עקרון: כל lag מחושב רק מהיסטוריה שקדמה לכל שורה — קפדנות מוחלטת על מניעת leakage.")

# ----- 6. Methodology -----
doc.add_page_break()
add_heading(doc, "6. מתודולוגיה — מניעת Leakage", level=1)

add_heading(doc, "12 עמודות שהוצאו במפורש", level=2)
add_para(doc, "נמדדות אחרי השידור: נתח, צופים 4+, חשיפה 4+, משך צפייה")
add_para(doc, "רייטינג מתחרים בזמן אמת: כאן 11, קשת 12, רשת 13, עכשיו 14")
add_para(doc, "נגזרות: ממוצע מתחרים, יתרון מול מתחרים, HUT proxy")
add_para(doc, "נגזרת ישירה מ-target: רייטינג מותאם")

add_heading(doc, "פיצול Train/Test", level=2)
add_para(doc, "כרונולוגי 80/20 (חיתוך 2026-02-08), לא random.")
add_callout(doc, "אם הייתי עושה random split היו לי R²=0.85 שווא, ובמציאות המודל היה נכשל.")

# ----- 7. Models -----
add_heading(doc, "7. המודלים שאומנו — 19 מודלים מ-7 משפחות", level=1)
add_table(doc,
    headers=["משפחה", "מודלים"],
    rows=[
        ("Baseline", "Naive Global Mean, Slot Mean"),
        ("ליניאריים", "Ridge, Lasso, ElasticNet, BayesianRidge, HuberRegressor"),
        ("מרחק/קרנל", "KNN, SVR (RBF)"),
        ("עץ בודד", "DecisionTree"),
        ("Tree Ensembles", "RF, ExtraTrees, GB, HistGB, XGBoost, LightGBM, CatBoost"),
        ("רשת נוירונים", "MLP (64×32)"),
        ("Stacking", "Ridge meta over RF+XGB+LGB"),
    ],
)

# ----- 8. Results -----
doc.add_page_break()
add_heading(doc, "8. תוצאות — Leaderboard", level=1)
add_table(doc,
    headers=["מקום", "מודל", "MAE", "R²"],
    rows=[
        ("🥇", "HistGradientBoosting", "0.263", "0.603"),
        ("🥈", "LightGBM", "0.265", "0.598"),
        ("🥉", "GradientBoosting", "0.270", "0.579"),
        ("4", "CatBoost", "0.271", "0.576"),
        ("5", "Stacking (Ridge meta)", "0.272", "0.566"),
        ("6", "ExtraTrees", "0.272", "0.579"),
        ("7", "XGBoost", "0.280", "0.558"),
        ("8", "RandomForest tuned", "0.280", "0.566"),
        ("9", "Lasso", "0.288", "0.486"),
        ("10", "ElasticNet", "0.292", "0.546"),
        ("11", "KNN", "0.294", "0.483"),
        ("12", "DecisionTree", "0.294", "0.467"),
        ("13", "SVR (RBF)", "0.298", "0.472"),
        ("14", "Slot Mean baseline", "0.305", "0.435"),
        ("15", "MLP", "0.328", "0.468"),
        ("16", "HuberRegressor", "0.343", "0.499"),
        ("17", "BayesianRidge", "0.355", "0.501"),
        ("18", "Ridge", "0.372", "0.471"),
        ("19", "Naive Global Mean", "0.422", "-0.046"),
    ],
)
add_callout(doc, "ההפרש בין מקום 1 ל-2 = 0.002 (רעש). השיפור על הנאיבי = 37.8% — שם הערך האמיתי של ML.")

# ----- 9. Error Analysis -----
doc.add_page_break()
add_heading(doc, "9. ניתוח שגיאות לפי מודל", level=1)

add_heading(doc, "9.1 דפוסים משותפים", level=2)
add_bullets(doc, [
    "שבת — הקשה ביותר אצל כולם (MAE 0.38-0.48)",
    "פריים-טיים 18-21 — חלק היום עם השגיאה הגבוהה ביותר (MAE 0.40-0.54)",
    "אירועים ביטחוניים — MAE 0.74-0.94 (פי 3 מהשגרה!) — תקרה אמיתית",
])

add_heading(doc, "9.2 פסקה לכל מודל (לראיון)", level=2)

per_model = [
    ("🥇 HistGradientBoosting — המנצח",
     "boosting עם binning של ה-features ל-256 דליים. ה-binning הוא רגולריזציה אוטומטית — "
     "מגביל את הרזולוציה של הפיצולים, לכן פחות overfit על דאטה קטן (~8K שורות). "
     "השיעור: על דאטה טבלאי קטן-בינוני, histogram-based GBDT הוא state-of-the-art."),

    ("🥈 LightGBM — שני בהפרש זעיר",
     "משפחה אחת עם HistGB אבל עם leaf-wise growth. ההפרש טכני-מימוש, לא קונצפטואלי."),

    ("🥉 GradientBoosting — איטי אבל איכותי",
     "Boosting קלאסי. MAE 0.270 מצוין אבל לקח 13.5 שניות (פי 12 מ-HistGB)."),

    ("CatBoost — היתרון לא בא לידי ביטוי",
     "boosting עם טיפול מובנה בקטגוריאליים. עשיתי one-hot ידני קודם — היתרון מקוזז. "
     "שיעור על coupling."),

    ("Stacking — לא עזר, ולמה זה בסדר",
     "meta-מודל שלמד לשקלל את חיזויי RF+XGB+LGB. MAE 0.272 — לא יותר טוב מ-LGB יחיד. "
     "Stacking עוזר רק כש-base models טועים בדברים שונים. אצלי כולם נכשלים באותם אירועים."),

    ("ExtraTrees — אקראיות נוספת עזרה",
     "כמו RF אבל הפיצול אקראי במקום אופטימלי. ניצח את RF (0.272 מול 0.280). "
     "שיעור: לפעמים פחות 'חכם' = יותר מדויק."),

    ("XGBoost — לא ניצח כצפוי",
     "אגדת קאגל. MAE 0.280, פחות טוב מ-LGB ו-HistGB. ערכים רציפים מזמינים overfitting על noise."),

    ("RandomForest — סוס העבודה",
     "400 עצים, bagging, max_features='sqrt'. MAE 0.280. בלי זמן לכוונון, RF מוציא 90%."),

    ("Lasso — ניצח את Ridge בזכות feature selection",
     "L1 מאפסת מקדמים מיותרים. MAE 0.288 — 0.084 יותר טוב מ-Ridge. "
     "אחרי one-hot על 6 קטגוריות יש 50+ features, רובם רעש."),

    ("ElasticNet — פשרה L1+L2",
     "MAE 0.292. כשFeatures מתואמים — זה מה שצריך."),

    ("KNN — נפל על מטריקת המרחק",
     "MAE 0.294. המרחק על one-hot של קטגוריות מטעה: שבת ו-ראשון נמצאים על מרחק יחידתי זהה."),

    ("DecisionTree — מראה מה ensemble מוסיף",
     "MAE 0.294. ההפרש בין 0.294 (עץ בודד) ל-0.263 (HistGB) הוא ה-bias-variance tradeoff."),

    ("SVR — kernel גמיש לא הספיק",
     "MAE 0.298. SVR לא מטפל טוב בקטגוריאליים גם עם kernel."),

    ("Slot Mean — Baseline חכם",
     "ממוצע (יום × שעה × שידור-חוזר). מסביר 43% מהשונות מ-feature אחד. "
     "מה שמומחה תוכן יעשה בלי ML."),

    ("MLP — נכשל בצפוי",
     "MAE 0.328. דאטה קטן + features מעורבים = MLP overfits. "
     "לא כל בעיה צריכה Deep Learning."),

    ("HuberRegressor — נכשל מסיבה מעניינת",
     "עמיד ל-outliers. חשבתי שיעזור באירועי-קיצון — אבל הוא מתעלם מ-outliers. "
     "אצלי outliers הם signal, לא noise."),

    ("Ridge, BayesianRidge — בודקים אם הקשר ליניארי",
     "נכשלים בצפוי (MAE 0.355-0.372). אם Ridge היה מנצח עצים — היה data leakage."),

    ("Naive Global Mean — הרצפה",
     "MAE 0.422. R² שלילי = drift חזק בין train ל-test."),
]

for name, desc in per_model:
    add_heading(doc, name, level=3)
    add_para(doc, desc)

# ----- 10. Performance Ceiling -----
doc.add_page_break()
add_heading(doc, "10. תקרת הביצועים", level=1)
add_heading(doc, "עשר השגיאות הגדולות ביותר — כולן באירועים ביטחוניים", level=2)
add_table(doc,
    headers=["תאריך", "תוכנית", "אמיתי", "חזוי", "שגיאה"],
    rows=[
        ("28/02/2026", "מיוחד-מתקפה באיראן", "5.02", "1.05", "3.97"),
        ("28/02/2026", "מיוחד-מתקפה באיראן", "4.69", "0.98", "3.71"),
        ("07/04/2026", "המהדורה המרכזית", "3.94", "0.78", "3.16"),
        ("03/04/2026", "קבינט שישי", "5.40", "2.33", "3.07"),
    ],
)
add_callout(doc, "כשמבצע צבאי מתחיל ב-28/02 ב-09:00, אין שום signal היסטורי שיגיד למודל "
                  "שב-12:00 אותו יום הרייטינג יהיה 5x מהממוצע. זאת התקרה האמיתית.")

# ----- 11. The Application -----
doc.add_page_break()
add_heading(doc, "11. האפליקציה — Streamlit חי", level=1)

add_heading(doc, "קישורים חיים", level=2)
add_bullets(doc, [
    "GitHub: https://github.com/DataScienceOritAlma/i24-ratings-forecast",
    "Streamlit Cloud: https://i24-ratings-orit.streamlit.app",
    "סיסמה: i24-2026-orit",
])

add_heading(doc, "5 מסכים פעילים", level=2)
add_bullets(doc, [
    "🏠 דף הבית — KPIs, leaderboard, ניווט",
    "📊 חיזויים — דפדוף וסינון של 1,957 חיזויים על מבחן",
    "📺 כרטיס תוכנית — drill-down ל-58 תוכניות עם חיזויים",
    "🔍 השוואת מודלים — leaderboard ויזואלי + MAE לפי חתך",
    "🎯 חיזוי עתידי — לאנליסט מחקר, עם רווח-בטחון וצופים מוערכים",
])

# ----- 12. Tech Stack -----
add_heading(doc, "12. Stack טכני", level=1)
add_table(doc,
    headers=["שכבה", "טכנולוגיה"],
    rows=[
        ("שפה", "Python 3.11"),
        ("נתונים", "pandas, numpy, openpyxl"),
        ("מודלינג", "scikit-learn 1.6, LightGBM, CatBoost, XGBoost"),
        ("Frontend", "Streamlit 1.40"),
        ("Visualization", "Plotly 5.24"),
        ("Auth", "password gate via Streamlit Secrets"),
        ("Hosting", "Streamlit Community Cloud"),
        ("Version control", "Git + GitHub"),
    ],
)

# ----- 13. Lessons -----
doc.add_page_break()
add_heading(doc, "13. לקחים — 7 תובנות מקצועיות", level=1)
lessons = [
    ("13.1 Baseline ראשון, מודל מתוחכם שני",
     "הנאיבי שלי הראה drift לפני שאיבדתי שעות ב-GridSearch."),
    ("13.2 לא כל אגדה מתאימה לבעיה",
     "XGBoost הוא אגדה, אבל פה HistGradientBoosting ניצח. דאטה קטן ו-binning עוזר."),
    ("13.3 outliers הם signal, לא noise",
     "HuberRegressor היה מהגרועים. באירועי-קיצון אני רוצה ללמוד מהם, לא להחליק."),
    ("13.4 Stacking לא קסם",
     "דורש diversity ב-errors. אם כולם נכשלים באותו אירוע — אין מה לשלב."),
    ("13.5 דע מתי להפסיק",
     "אחרי 3 גרסאות הבנתי שה-MAE הבא יבוא מ-features חדשים, לא ממודל."),
    ("13.6 חיזוי קדימה ≠ MAE על מבחן",
     "נדרשה תיקון טרנד וחלון 90 יום ל-lag features."),
    ("13.7 UX קובע ערך",
     "מודל מצוין בלי UX ראוי = חסר ערך."),
]
for title_t, desc in lessons:
    add_heading(doc, title_t, level=3)
    add_para(doc, desc)

# ----- 14. Future Work -----
add_heading(doc, "14. עבודה עתידית", level=1)
add_heading(doc, "Quick wins (חודש)", level=2)
add_bullets(doc, [
    "Quantile regression — חיזוי טווח במקום נקודה",
    "Cold-start handling — שיטה ייעודית לתוכניות בלי היסטוריה",
    "מודל יומי מצטבר על-בסיס המודל הקיים",
])
add_heading(doc, "Mid-term (3 חודשים)", level=2)
add_bullets(doc, [
    "חיבור ל-feed דיגיטלי של i24 — ייתכן ינבא breaking",
    "A/B test של החיזויים מול מנהלי תוכן",
    "Embedding לשמות תוכניות (text features)",
])
add_heading(doc, "Long-term (6+ חודשים)", level=2)
add_bullets(doc, [
    "Multi-target: רייטינג + נתח + צפייה דמוגרפית",
    "Real-time updating על-בסיס social signals",
    "מוצר B2B — להציע למתחרים",
])

# ----- 15. Interview Q&A -----
doc.add_page_break()
add_heading(doc, "15. נספח: שאלות ראיון", level=1)

qa = [
    ("למה לא רשת נוירונים?",
     "ניסיתי MLP, MAE=0.328 — מהגרועים. דאטה של 10K שורות קטן מדי לרשתות. "
     "על דאטה טבלאי קטן GBDT מנצח כל פעם."),
    ("איך אתה יודע שאין leakage?",
     "12 עמודות הוצאו במפורש (post-airing + competitor real-time). "
     "lag features מחושבים cumulative. פיצול כרונולוגי."),
    ("למה לא רגרסיה זמנית קלאסית?",
     "ניסיתי SARIMAX ו-Additive. הם נכשלו עם R² שלילי כי הם גלובליים-יומיים."),
    ("מה אם יהיו לך 100K שורות?",
     "כדאי לנסות שוב MLP/Transformer. אבל הצפי שלי הוא ש-LightGBM יישאר תחרותי."),
    ("איך אתה מטפל בחיזוי קדימה?",
     "שתי שכבות תיקון: חלון 90 ימים ל-lag features, ותיקון טרנד "
     "מבוסס linear regression על 6 חודשים, capped ±5%/חודש."),
]
for q, a in qa:
    add_heading(doc, f"ש: {q}", level=3)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ת: ")
    run.font.bold = True
    run.font.name = "David"
    run.font.size = Pt(11)
    run.font.color.rgb = PRIMARY
    run = p.add_run(a)
    run.font.name = "David"
    run.font.size = Pt(11)

# ----- Summary table -----
add_heading(doc, "סיכום מספרי", level=1)
add_table(doc,
    headers=["מטריקה", "ערך"],
    rows=[
        ("שורות בדאטה", "10,039"),
        ("תוכניות ייחודיות", "179"),
        ("מודלים שנוסו", "19"),
        ("משפחות מודלים", "7"),
        ("MAE של המנצח", "0.263"),
        ("R² של המנצח", "0.603"),
        ("שיפור על הנאיבי", "37.8%"),
        ("מסכי האפליקציה", "5"),
        ("Commits ל-GitHub", "17"),
        ("קבצי תיעוד", "10"),
    ],
)

# ----- Final quote -----
doc.add_paragraph()
add_quote(doc,
    "זה לא היה רק פרוייקט סיום קורס. זה היה תהליך של 12 חודשים שבו עברתי מ-EDA, "
    "דרך 3 גרסאות מידול, השוואת 19 מודלים, ניתוח שגיאות מעמיק, "
    "ועד מערכת חיה עם URL לשיתוף. כל החלטה נשענת על מספר וכל מספר נשען על קוד שאפשר להריץ מחדש."
)
add_para(doc, "— אורית עלמה זיו-נר, מאי 2026", align="left", color=MUTED)

# Save
output = "פרוייקט_i24_סיכום.docx"
doc.save(output)
print(f"Saved: {output}")
import os
print(f"Size: {os.path.getsize(output) / 1024:.0f} KB")
